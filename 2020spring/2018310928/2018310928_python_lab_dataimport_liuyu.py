#1
with open("data/BABAnews.txt", mode = 'r',encoding="utf-8") as myfile:
    print(myfile.read())
    myfile.closed

#2
myfile = open("data/BABAnews.txt", mode = 'r',encoding="utf-8")
text = myfile.read()
myfile.close()
print(text)

#3
myfile = open("data/BABAnews.txt", mode = 'r',encoding="utf-8")
print(myfile.read())
print(myfile.readline())

#4
with open('data/Titanic.csv', 'r',encoding="utf-8") as myfile:
    for line in myfile:
        print(line, end ='')

#5
file = open('data/testfile.txt', 'w') 
file.write('Hello World \n') 
file.write('and this is another line.') 
file.close()

#6
tobuy = ['milk\n', 'butter\n', 'coffee beans\n', 'arugula\n']
file = open('data/grocerylist.txt', 'w')
file.writelines(tobuy) # writelines(list)
file.close()

#7
pi = 3.141592
fout = open('data/math.txt', 'w')
fout.write("Pi's value is "+str(pi))

#8
import numpy as np 
data = np.loadtxt('data/BJsales.txt', skiprows = 1, delimiter='\t')
print(data)

#9
import csv
with open('data/Titanic.csv') as csvfile:
    titanicReader = csv.reader(csvfile)
    for row in titanicReader:
        print('  '.join(row))

#10
import pandas as pd
titanicData = pd.read_csv('data/Titanic.csv') 
titanicData.head()

#11
titanicData = pd.read_table('data/Titanic.csv', sep=',')
titanicData.head()

#12
import pandas as pd
file = 'data/example.xlsx'
myfile = pd.ExcelFile(file) 
print(myfile.sheet_names) 
dataframe = myfile.parse('HairEyeColor') 
dataframe.head()

#13
import PyPDF2
pdfFileObj = open('data/Li2011Wiley.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())
pdfFileObj.close()

#14
import docx
doc = docx.Document('data/Li2011Wiley.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)

#15
import requests
url = 'https://sp0.baidu.com/8aQDcjqpAAV3otqbppnN2DJv/api.php?resource_id=6899&query=失信执行人名单&iname=中国银行'
resp = requests.get(url)

import json
data = json.loads(resp.text, encoding='gb18030')
data.keys()

#16
result = data['data'][0]['result']
for i in range(len(result)):
        data = result[i]
        caseCode = data['caseCode']
        areaName = data['areaName']
        businessEntity = data['businessEntity']
        courtName = data['courtName']
        duty = data['duty']
        performance = data['performance']
        disruptTypeName = data['disruptTypeName']
        publishDate = data['publishDate']
        regDate = data['regDate']
        gistId = data['gistId']
        gistUnit = data['gistUnit']
        cardNum = data['cardNum']
        print('  '.join([caseCode, areaName, businessEntity, courtName, duty, performance, disruptTypeName, publishDate, regDate, gistId, gistUnit, cardNum]))

#17
import sqlite3
connection = sqlite3.connect("survey.db")
cursor = connection.cursor()
cursor.execute("SELECT Site.lat, Site.long FROM Site;")
results = cursor.fetchall()
for r in results:
    print(r)
cursor.close()
connection.close()



















