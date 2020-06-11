#!/usr/bin/env python
# coding: utf-8

# # Data import and storage
# Input and output typically falls into a few main categories: reading text files and other more efficient on-disk formats like pdf and docx files, interacting with net-work sources like web APIs, and loading data from databases.

# ## Reading and writing a text file
# ### Using `open()` to read and write

# `open()` returns a file object, and is most commonly used with two arguments: `open(filename, mode)`. It is good practice to use the with keyword when dealing with file objects. The advantage is that the file is properly closed after its suite finishes.
# 
# - `mode` can be 'r' for read
# - 'w' for only writing
# - 'a' opens the file for appending
# - 'r+' opens the file for both reading and writing
# 

# In[3]:


with open("data/BABAnews.txt", mode = 'r') as myfile:
    print(myfile.read())


# In[5]:


with open("data/BABAnews.txt", mode = 'r',encoding='utf-8') as myfile:
    print(myfile.read())


# In[6]:


myfile.closed


# In[37]:


with open("data/BABAnews.txt", mode = 'r') as myfile: 
        print(myfile.read())


# In[38]:


myfile.closed


# If you’re not using the with keyword, then you should call f.close() to close the file and immediately free up any system resources used by it. For example:

# In[9]:


myfile=open("data/BABAnews.txt",mode='r',encoding='utf-8')
text=myfile.read()
myfile.close()
print(text)


# In[9]:


myfile = open("data/BABAnews.txt",mode='r',encoding='utf-8')
text = myfile.read()
myfile.close()
print(text)


# In[72]:


myfile = open("data/BABAnews.txt", mode = 'r')
text = myfile.read()
myfile.close()
print(text)


# In[10]:


myfile.read()


# In[10]:


myfile.read()


# In[42]:


myfile.read() # can be read after the file being closed


# To read the contents of a file, we use `f.read(size)`. 

# In[6]:


myfile=open("data/BABAnews.txt",mode='r',encoding='utf-8')
myfile.read()


# In[4]:


myfile.read()
myfile.close()


# In[48]:


myfile = open("data/BABAnews.txt", mode = 'r')
myfile.read()


# In[49]:


myfile.read()
myfile.close()


# `f.readline()` reads a single line from the file; a newline character (\n) is left at the end of the string, and is only omitted on the last line of the file if the file doesn’t end in a newline. This makes the return value unambiguous; if `f.readline()` returns an empty string, the end of the file has been reached, while a blank line is represented by '\n', a string containing only a single newline. 
# 
# If you want to read all the lines of a file in a list you can also use `list(f)` or `f.readlines()`.

# In[77]:


myfile = open("data/BABAnews.txt", mode = 'r')
myfile.readline()


# In[78]:


myfile.readlines()


# **Note:** Both `read()` and `readlines()` come with the concept of a cursor. After either command is executed, the cursor moves to the end of the file, leaving nothing more to read in. Therefore, once a file content has been read in, another attempt to read from the file object will produce an empty data object. If for some reason you must read the file content again, you must close and re-open the file. 
# 
# Lastly, rather than loading the entire file content into memory, you can iterate through the file object line by line using the `for` loop. This method is more memory-efficient and therefore recommended when dealing with a very large file. 

# In[16]:


with open('data/Titanic.csv','r') as myfile:
    for line in myfile:
        print(line, end='')


# In[13]:


with open('data/Titanic.csv', 'r') as myfile:
    for line in myfile:
        print(line, end ='')


# In[12]:


with open('data/Titanic.csv', 'r') as myfile:
    for line in myfile:
        print(line, end ='')


# Writing methods also come in a pair: `write()` and `writelines()`. Like the corresponding reading methods, `write()` handles a single string, while `writelines()` handles a list of strings. 
# 
# Below, `write()` writes a single string each time to the designated output file:

# In[20]:


file=open('data/testfile.txt','w')
file.write('Hello World \n')
file.write('and this is another line.')
file.close()


# In[66]:


file = open('data/testfile.txt', 'w') 
file.write('Hello World \n') 
file.write('and this is another line.') 
file.close() 


# In[ ]:


This time, we have `tobuy`, a list of strings, which `writelines()` writes out at once:


# In[24]:


tobuy=['milk\n','butter\n','coffee beans\n','arugula\n']
file=open('data/grocerylist.txt','w')
file.writelines(tobuy)
file.close()


# In[79]:


tobuy = ['milk\n', 'butter\n', 'coffee beans\n', 'arugula\n']
file = open('data/grocerylist.txt', 'w')
file.writelines(tobuy) # writelines(list)
file.close()


# **Only the string type can be written.** Writing methods only works with strings: `write()` takes a single string, and `writelines()` takes a list which contains strings only. Non-string type data must be first coerced into the string type by using the `str()` function. 

# In[28]:


pi=3.141592
fout=open('data/math.txt','w')
fout.write("Pi's value is")
fout.write(str(pi))
fout.close()


# In[80]:


pi = 3.141592
fout = open('data/math.txt', 'w')
fout.write("Pi's value is ")
fout.write(pi) # trying to write float, doesn't work


# In[81]:


fout.write(str(pi))
fout.close()


# ### Other ways to read text files

# We can also use other modules to read text files. For example, we can use **numpy** to read *txt* file.  **csv** and **pandas** can be used to read *csv* files.

# In[29]:


import numpy as np
data=np.loadtxt('data/BJsales.txt',skiprows=1,delimiter='\t')
print(data)


# In[88]:


import numpy as np 
data = np.loadtxt('data/BJsales.txt', skiprows = 1, delimiter='\t')
print(data)


# In[60]:


import csv
with open('data/Titanic.csv') as csvfile:
    titanicReader = csv.reader(csvfile)
    for row in titanicReader:
        print('  '.join(row))


# The `pandas` I/O API is a set of top level `reader` functions accessed like `read_csv()` that generally return a pandas object. These functions includes
# 
#     read_excel
#     read_hdf
#     read_sql
#     read_json
#     read_msgpack (experimental)
#     read_html
#     read_gbq (experimental)
#     read_stata
#     read_sas
#     read_clipboard
#     read_pickle
#     
# See [pandas IO tools](http://pandas.pydata.org/pandas-docs/stable/io.html) for detailed explanation.

# In[34]:


import pandas as pd
titanicData=pd.read_csv('data/Titanic.csv')
titanicData.head()


# In[33]:


import pandas as pd
titanicData = pd.read_csv('data/Titanic.csv') 
titanicData.head()


# In[61]:


import pandas as pd
titanicData = pd.read_csv('data/Titanic.csv') 
titanicData.head()


# We could also have used `read_table()` and specifying the delimiter:

# In[35]:


titanicData=pd.read_table('data/Titanic.csv',sep=',')
titanicData.head()


# In[84]:


titanicData = pd.read_table('data/Titanic.csv', sep=',')
titanicData.head()


# A file will not always have a header row. To read this in, you have a couple of options. You can allow pandas to assign default column names, or you can specify names yourself:

# In[36]:


sales=pd.read_csv('data/BJsales.csv',header=None)
sales.head()


# In[92]:


sales = pd.read_csv('data/BJsales.csv', header = None)
sales.head()


# In[37]:


sales = pd.read_csv('data/BJsales.csv', names=['ID','Time','Value'])
sales.head()


# In[40]:


sales = pd.read_csv('data/BJsales.csv', names=[ 'ID','Time','Value'])
sales.head()


# Data can also be exported to delimited format. Let’s consider one of the CSV files read above:

# In[41]:


sales.to_csv('data/sales.csv', index = False)


# ## Reading a Microsoft Excel file using `pandas`

# Excel files are a huge part of any business operation and it becomes imperative that you learn exactly how to import these into python for data analysis. In order to do this we can use the code snippet shown below:

# In[45]:


import pandas as pd
file='data/example.xlsx'
myfile=pd.ExcelFile(file)
print(myfile.sheet_names)
dataframe=myfile.parse('HairEyeColor')
dataframe.head()


# In[44]:


import pandas as pd
file = 'data/example.xlsx'
myfile = pd.ExcelFile(file) 
print(myfile.sheet_names) 
dataframe = myfile.parse('HairEyeColor') 
dataframe.head()


# In[102]:


import pandas as pd
file = 'data/example.xlsx'
myfile = pd.ExcelFile(file) 
print(myfile.sheet_names) #printing out all the sheet names in the excel file
dataframe = myfile.parse('HairEyeColor') #extracting data from the first sheet as a dataframe
dataframe.head()


# ## Reading a pdf file using `PyPDF2`
# 
# `PyPDF2` is a python library built as a PDF toolkit. It is capable of:
# 
# - Extracting document information (title, author, …)
# - Splitting documents page by page
# - Merging documents page by page
# - Cropping pages
# - Merging multiple pages into a single page
# - Encrypting and decrypting PDF files
# - and more!
# 
# Here we only demonstrate how to read a pdf file. Please find more details on https://pypi.org/project/PyPDF2/.

# In[47]:


import PyPDF2
pdfFileObj = open('data/Li2011Wiley.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())
pdfFileObj.close()


# In[8]:


import PyPDF2
pdfFileObj = open('data/Li2011Wiley.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())
pdfFileObj.close()


# ## Reading a word file using `docx`

# In[49]:


import docx
doc=docx.Document('data/Li2011Wiley.docx')
print(len(doc.paragraphs))
print(doc.paragrarphs[0].text)


# In[50]:


import docx
doc = docx.Document('data/Li2011Wiley.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)


# In[127]:


import docx
doc = docx.Document('data/Li2011Wiley.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)


# ## Interacting with HTML and Web APIs

# Many websites have public APIs providing data feeds via JSON or some other format. There are a number of ways to access these APIs from Python; one easy-to-use method that I recommend is the requests package (http://docs.python-requests.org). 

# In[52]:


import requests
url='https://sp0.baidu.com/8aQDcjqpAAV3otqbppnN2DJv/api.php?resource_id=6899&query=失信执行人名单&iname=中国银行'
resp=requests.get(url)
import json
data=json.loads(resp.text,encoding='gb18030')
data.keys()


# In[53]:


import requests
url = 'https://sp0.baidu.com/8aQDcjqpAAV3otqbppnN2DJv/api.php?resource_id=6899&query=失信执行人名单&iname=中国银行'
resp = requests.get(url)
import json
data = json.loads(resp.text, encoding='gb18030')
data.keys()


# In[136]:


import requests
url = 'https://sp0.baidu.com/8aQDcjqpAAV3otqbppnN2DJv/api.php?resource_id=6899&query=失信执行人名单&iname=中国银行'
resp = requests.get(url)
import json
data = json.loads(resp.text, encoding='gb18030')
data.keys()


# In[54]:


result = data['data'][0]['result']
for i in range(len(result)):
        data = result[i]
        caseCode = data['caseCode']
        areaName = data['areaName']
        businessEntity = data['businessEntity']
        courtName = data['courtName']
        duty = data['duty']
        performance = data['performance']
        disruptTypeName = data['disruptTypeName']
        publishDate = data['publishDate']
        regDate = data['regDate']
        gistId = data['gistId']
        gistUnit = data['gistUnit']
        cardNum = data['cardNum']
        print('  '.join([caseCode, areaName, businessEntity, courtName, duty, performance, disruptTypeName, publishDate, regDate, gistId, gistUnit, cardNum]))


# In[137]:


result = data['data'][0]['result']
for i in range(len(result)):
        data = result[i]
        caseCode = data['caseCode']
        areaName = data['areaName']
        businessEntity = data['businessEntity']
        courtName = data['courtName']
        duty = data['duty']
        performance = data['performance']
        disruptTypeName = data['disruptTypeName']
        publishDate = data['publishDate']
        regDate = data['regDate']
        gistId = data['gistId']
        gistUnit = data['gistUnit']
        cardNum = data['cardNum']
        print('  '.join([caseCode, areaName, businessEntity, courtName, duty, performance, disruptTypeName, publishDate, regDate, gistId, gistUnit, cardNum]))


# ## Interacting with Databases

# In many applications data rarely comes from text files, that being a fairly inefficient way to store large amounts of data. SQL-based relational databases (such as SQL Server, PostgreSQL, and MySQL) are in wide use, and many alternative non-SQL (so-called NoSQL) databases have become quite popular. The choice of database is usually de- pendent on the performance, data integrity, and scalability needs of an application.
# 
# Loading data from SQL into a DataFrame is fairly straightforward, and pandas has some functions to simplify the process. As an example, I’ll use an in-memory SQLite database using Python’s built-in sqlite3 driver. Here’s a short Python program that selects latitudes and longitudes from an SQLite database stored in a file called survey.db:

# In[57]:


import sqlite3
connection=sqlite3.connect("survey.db")
cursor=connection.cursor()
cursor.execute("SELECT Site.lat,Site.long FROM Site;")
results=cursor.fetchall()
for r in results:
    print(r)
cursor.close()
connection.close()


# In[ ]:


import sqlite3
connection = sqlite3.connect("survey.db")
cursor = connection.cursor()
cursor.execute("SELECT Site.lat, Site.long FROM Site;")
results = cursor.fetchall()
for r in results:
    print(r)
cursor.close()
connection.close()


# The program starts by importing the sqlite3 library. If we were connecting to MySQL, DB2, or some other database, we would import a different library, but all of them provide the same functions, so that the rest of our program does not have to change (at least, not much) if we switch from one database to another.
# 
# Line 2 establishes a connection to the database. Since we’re using SQLite, all we need to specify is the name of the database file. Other systems may require us to provide a username and password as well. Line 3 then uses this connection to create a cursor. Just like the cursor in an editor, its role is to keep track of where we are in the database.
# 
# On line 4, we use that cursor to ask the database to execute a query for us. The query is written in SQL, and passed to cursor.execute as a string. It’s our job to make sure that SQL is properly formatted; if it isn’t, or if something goes wrong when it is being executed, the database will report an error.
# 
# The database returns the results of the query to us in response to the cursor.fetchall call on line 5. This result is a list with one entry for each record in the result set; if we loop over that list (line 6) and print those list entries (line 7), we can see that each one is a tuple with one element for each field we asked for.
# 
# Finally, lines 8 and 9 close our cursor and our connection, since the database can only keep a limited number of these open at one time. Since establishing a connection takes time, though, we shouldn’t open a connection, do one operation, then close the connection, only to reopen it a few microseconds later to do another operation. Instead, it’s normal to create one connection that stays open for the lifetime of the program.

# # Lab
# 
# Read your five different types of your own files to Python from your hard disk.
