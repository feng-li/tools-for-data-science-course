#!/usr/bin/env python
# coding: utf-8

# # Data import and storage
# Input and output typically falls into a few main categories: reading text files and other more efficient on-disk formats like pdf and docx files, interacting with net-work sources like web APIs, and loading data from databases.

# ## Reading and writing a text file
# ### Using `open()` to read and write

# `open()` returns a file object, and is most commonly used with two arguments: `open(filename, mode)`. It is good practice to use the with keyword when dealing with file objects. The advantage is that the file is properly closed after its suite finishes.
# 
# - `mode` can be 'r' for read
# - 'w' for only writing
# - 'a' opens the file for appending
# - 'r+' opens the file for both reading and writing
# 

# In[13]:


with open("C:\\Users\\wendy.wu\\Desktop\\数据科学工具\\python\\data practice\\BABAnews.txt", encoding="utf-8",mode = 'r') as myfile: 
        print(myfile.read())


# In[14]:


myfile.closed


# If you’re not using the with keyword, then you should call f.close() to close the file and immediately free up any system resources used by it. For example:

# In[17]:


myfile = open("C:\\Users\\wendy.wu\\Desktop\\数据科学工具\\python\\data practice\\BABAnews.txt", encoding="utf-8",mode = 'r')
text = myfile.read()
myfile.close()
print(text)


# In[22]:


myfile.read() # can't be read after the file being closed


# To read the contents of a file, we use `f.read(size)`. 

# In[35]:


myfile = open("C:\\Users\\wendy.wu\\Desktop\\数据科学工具\\python\\data practice\\BABAnews.txt", encoding="utf-8", mode = 'r')
myfile.read()


# In[36]:


myfile.read()


# In[37]:


myfile.close()


# `f.readline()` reads a single line from the file; a newline character (\n) is left at the end of the string, and is only omitted（省去） on the last line of the file if the file doesn’t end in a newline. This makes the return value unambiguous; if `f.readline()` returns an empty string, the end of the file has been reached, while a blank line is represented by '\n', a string containing only a single newline. 
# 
# If you want to read all the lines of a file in a list you can also use `list(f)` or `f.readlines()`.

# In[39]:


myfile = open("C:\\Users\\wendy.wu\\Desktop\\数据科学工具\\python\\data practice\\BABAnews.txt", encoding="utf-8", mode = 'r')
myfile.readline()


# In[40]:


myfile.readlines()


# **Note:** Both `read()` and `readlines()` come with the concept of a cursor（光标）. After either command is executed, the cursor moves to *the end of the file*, leaving nothing more to read in. Therefore, once a file content has been read in, another attempt to read from the file object will produce an empty data object. If for some reason you must read the file content again, you must close and re-open the file. 
# 
# Lastly, rather than loading the entire file content into memory, you can iterate through the file object *line by line* using the `for` loop. This method is more memory-efficient and therefore ***recommended*** when dealing with a very large file. 

# In[54]:


with open("C:\\Users\\wendy.wu\\Desktop\\数据科学工具\\python\\data practice\\Titanic.csv",  'r') as myfile:
    for line in myfile:
        print(line, end ='*')


# Writing methods also come in a pair: `write()` and `writelines()`. Like the corresponding reading methods, `write()` handles a single string, while `writelines()` handles a list of strings. 
# 
# Below, `write()` writes a single string each time to the designated output file:

# In[50]:


file = open("C:\\Users\\wendy.wu\\Desktop\\数据科学工具\\python\\data practice\\testfile.txt", 'w') 
file.write('Hello World \n') 
file.write('and this is another line.') 
file.close() 


# In[61]:


with open("C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/testfile.txt", mode='r') as myfile:
   print(myfile.read())


# This time, we have `tobuy`, *a list of strings*, which `writelines()` writes out at once:

# In[55]:


tobuy = ['milk\n', 'butter\n', 'coffee beans\n', 'arugula\n']
file = open('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/grocerylist.txt', 'w')
file.writelines(tobuy) # writelines(list)
file.close()


# In[60]:


with open("C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/grocerylist.txt", mode='r') as myfile:
   print(myfile.read())


# **Only the string type can be *written*.** Writing methods only works with strings: `write()` takes a single string, and `writelines()` takes a list which contains strings only. Non-string type data must be first coerced（强迫） into the string type by using the `str()` function. 

# In[62]:


pi = 3.141592
fout = open('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/math.txt', 'w')
fout.write("Pi's value is ")
fout.write(pi) # trying to write float, doesn't work


# In[63]:


fout.write(str(pi))
fout.close()


# In[68]:


with open('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/math.txt', 'r') as myfile:
   print(myfile.read())


# ### Other ways to read text files

# We can also use other modules to read text files. For example, we can use **numpy** to read *txt* file.  **csv** and **pandas** can be used to read *csv* files.

# In[70]:


import numpy as np 
data = np.loadtxt('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/BJsales.txt', skiprows = 1, delimiter='\t')
print(data)


# In[71]:


import csv
with open('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/Titanic.csv') as csvfile:
    titanicReader = csv.reader(csvfile)
    for row in titanicReader:
        print('  '.join(row))


# The `pandas` I/O API is a set of top level `reader` functions accessed like `read_csv()` that generally return a pandas object. These functions includes
# 
#     read_excel
#     read_hdf
#     read_sql
#     read_json
#     read_msgpack (experimental)
#     read_html
#     read_gbq (experimental)
#     read_stata
#     read_sas
#     read_clipboard
#     read_pickle
#     
# See [pandas IO tools](http://pandas.pydata.org/pandas-docs/stable/io.html) for detailed explanation.

# In[73]:


import pandas as pd
titanicData = pd.read_csv('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/Titanic.csv') 
titanicData.head()


# We could also have used `read_table()` and specifying the delimiter（定界符）:

# In[84]:


titanicData = pd.read_table('data/Titanic.csv', sep=',')
titanicData.head()


# A file will not always have a header row. To read this in, you have a couple of options. You can allow pandas to assign default column names, or you can specify names yourself:

# In[74]:


sales = pd.read_csv('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/BJsales.csv', header = None)
sales.head()


# In[75]:


sales = pd.read_csv('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/BJsales.csv', names=['ID', 'Time', 'Value'])
sales.head()


# Data can also be exported to delimited format. Let’s consider one of the CSV files read above:

# In[79]:


sales.to_csv('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/sales.csv', index =False)
sales=pd.read_csv('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/sales.csv', names=['ID', 'Time', 'Value'])
sales.head()


# ## Reading a Microsoft Excel file using `pandas`

# Excel files are a huge part of any business operation and it becomes imperative that you learn exactly how to import these into python for data analysis. In order to do this we can use the code snippet shown below:

# In[1]:


import pandas as pd
file = 'C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/example.xlsx'
myfile = pd.ExcelFile(file) 
print(myfile.sheet_names) #printing out all the sheet names in the excel file
dataframe = myfile.parse('HairEyeColor') #extracting data from the first sheet as a dataframe
dataframe.head()


# ## Reading a pdf file using `PyPDF2`
# 
# `PyPDF2` is a python library built as a PDF toolkit. It is capable of:
# 
# - Extracting document information (title, author, …)
# - Splitting（分离） documents page by page
# - Merging documents page by page
# - Cropping（修剪） pages
# - Merging multiple pages into a single page
# - Encrypting（加密） and decrypting PDF files
# - and more!
# 
# Here we only demonstrate how to read a pdf file. Please find more details on https://pypi.org/project/PyPDF2/.

# In[12]:


get_ipython().system(' pip install PyPDF2')


# In[14]:


import PyPDF2
pdfFileObj = open('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/Li2011Wiley.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())
pdfFileObj.close()


# ## Reading a word file using `docx`

# In[15]:


get_ipython().system(' pip install docx')


# In[21]:


import docx
doc = docx.Document('C:/Users/wendy.wu/Desktop/数据科学工具/python/data practice/Li2011Wiley.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)


# ## Interacting with HTML and Web APIs

# Many websites have public APIs providing data feeds via JSON or some other format. There are a number of ways to access these APIs from Python; one easy-to-use method that I recommend is the requests package (http://docs.python-requests.org). 

# In[25]:


import requests
url = 'https://www.cnblogs.com/hui-code/p/12030353.html'
resp = requests.get(url)
import json
data = json.loads(resp.text,encoding='gb18030')
data.keys()


# In[24]:


result = data['data'][0]['result']
for i in range(len(result)):
        data = result[i]
        caseCode = data['caseCode']
        areaName = data['areaName']
        businessEntity = data['businessEntity']
        courtName = data['courtName']
        duty = data['duty']
        performance = data['performance']
        disruptTypeName = data['disruptTypeName']
        publishDate = data['publishDate']
        regDate = data['regDate']
        gistId = data['gistId']
        gistUnit = data['gistUnit']
        cardNum = data['cardNum']
        print('  '.join([caseCode, areaName, businessEntity, courtName, duty, performance, disruptTypeName, publishDate, regDate, gistId, gistUnit, cardNum]))


# ## Interacting with Databases

# In many applications data rarely comes from text files, that being a fairly inefficient way to store large amounts of data. **SQL-based** relational databases (such as SQL Server, PostgreSQL, and MySQL) are in wide use, and many alternative non-SQL (so-called NoSQL) databases have become quite popular. The choice of database is usually de- pendent on the performance, data integrity, and scalability needs of an application.
# 
# Loading data from SQL into a DataFrame is fairly straightforward, and pandas has some functions to simplify the process. As an example, I’ll use an in-memory SQLite database using Python’s built-in sqlite3 driver. Here’s a short Python program that selects latitudes and longitudes from an SQLite database stored in a file called survey.db:

# In[7]:


import sqlite3
connection = sqlite3.connect("survey.db")
cursor = connection.cursor()
cursor.execute("SELECT Site.lat, Site.long FROM Site;")
results = cursor.fetchall()
for r in results:
    print(r)
cursor.close()
connection.close()


# The program starts by importing the sqlite3 library. If we were connecting to MySQL, DB2, or some other database, we would import a different library, but all of them provide the same functions, so that the rest of our program does not have to change (at least, not much) if we switch from one database to another.
# 
# Line 2 establishes a connection to the database. Since we’re using SQLite, all we need to specify is the name of the database file. Other systems may require us to provide a username and password as well. Line 3 then uses this connection to create a cursor. Just like the cursor in an editor, its role is to keep track of where we are in the database.
# 
# On line 4, we use that cursor to ask the database to execute a query for us. The query is written in SQL, and passed to cursor.execute as a string. It’s our job to make sure that SQL is properly formatted; if it isn’t, or if something goes wrong when it is being executed, the database will report an error.
# 
# The database returns the results of the query to us in response to the cursor.fetchall call on line 5. This result is a list with one entry for each record in the result set; if we loop over that list (line 6) and print those list entries (line 7), we can see that each one is a tuple with one element for each field we asked for.
# 
# Finally, lines 8 and 9 close our cursor and our connection, since the database can only keep a limited number of these open at one time. Since establishing a connection takes time, though, we shouldn’t open a connection, do one operation, then close the connection, only to reopen it a few microseconds later to do another operation. Instead, it’s normal to create one connection that stays open for the lifetime of the program.

# # Lab
# 
# Read your five different types of your own files to Python from your hard disk.
