{
 "cells": [
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "# Data import and storage\n",
    "Input and output typically falls into a few main categories: reading text files and other more efficient on-disk formats like pdf and docx files, interacting with net-work sources like web APIs, and loading data from databases."
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "## Reading and writing a text file\n",
    "### Using `open()` to read and write"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "`open()` returns a file object, and is most commonly used with two arguments: `open(filename, mode)`. It is good practice to use the with keyword when dealing with file objects. The advantage is that the file is properly closed after its suite finishes.\n",
    "\n",
    "- `mode` can be 'r' for read\n",
    "- 'w' for only writing\n",
    "- 'a' opens the file for appending\n",
    "- 'r+' opens the file for both reading and writing\n"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 12,
   "metadata": {},
   "outputs": [
    {
     "data": {
      "image/png": "iVBORw0KGgoAAAANSUhEUgAAAXoAAAD4CAYAAADiry33AAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAALEgAACxIB0t1+/AAAADh0RVh0U29mdHdhcmUAbWF0cGxvdGxpYiB2ZXJzaW9uMy4xLjMsIGh0dHA6Ly9tYXRwbG90bGliLm9yZy+AADFEAAAgAElEQVR4nO3deZRV5ZX38e+mmGRQERCRGcQIKIqWA44oKCAKCGLANtFuE9skdLrbJK3dppO1tNMrb+wkncEkmsRMHUPACUQQUcEJFQpBBBEZIoIokwgoMhS13z+eS3K55xRcqm6dO/0+a9Wi7j7n3trUqtp17nOeZz/m7oiISOlqlO8ERESkYanQi4iUOBV6EZESp0IvIlLiVOhFREpc43wnkKldu3bevXv3fKchIlJUFi5cuMXd28cdK7hC3717d6qqqvKdhohIUTGztbUd09CNiEiJU6EXESlxKvQiIiVOhV5EpMRlVejNbJiZrTCzVWZ2xyHOu9bM3Mwq02L/nnreCjMbmoukRUQke4eddWNmFcC9wOXAemCBmU1z9zczzmsNfBV4NS3WFxgP9ANOBJ42s5PdfX/u/gsiDae6GjZvhu3b/xarqIC2baFNGzDLX24i2cpmeuU5wCp3XwNgZpOAUcCbGefdDXwP+HpabBQwyd33AH8xs1Wp13u5vomLNISdO+GNN8LHu+/C1q1QW4PXJk2gQwfo1Qv694fPfCbERApNNoW+E7Au7fF64Nz0E8xsANDF3aeb2dcznvtKxnM71TFXkQaxfz9UVcFzz8GaNbUX9kz79sH69eHjueegWbNQ8AcPhh49GjZnkSORTaGPe3P6118FM2sE/BC46Uifm/YatwC3AHTt2jWLlETqb88eeP55eOYZ2LYtN6+3YEH46NULhg4NhV/DO5Jv2RT69UCXtMedgQ1pj1sDpwJzLfxEnwBMM7ORWTwXAHe/H7gfoLKyUjuhSINyh4ULYcoU+Oijhvkaq1fDz34GJ58MEybAiSc2zNcRyUY2hX4B0NvMegDvEW6uXn/goLtvB9odeGxmc4Gvu3uVmX0KPGhmPyDcjO0NzM9d+iJHZuNGePBBeOut7J9z7LHQrl24CQuwe3d4nd27D//ct9+Gu++GSy+FUaPC8I5I0g5b6N292swmArOACuABd19mZncBVe4+7RDPXWZmkwk3bquBr2jGjeSDO8ybB3/6UxhbP5Tjj4fTTw/DLl27QvPm8a+3cyesXAmvvx5u3u7aFf96NTVheGjpUvjiF6FLl/jzRBqKFdqesZWVla6mZpJLu3eHq/hXX639HLNQ3C+/PIyvH+m4enV1GJufPRvee6/28xo3hnHj4JJLNHYvuWVmC929Mu5YwXWvFMmlLVvgJz+BDz6o/Zz+/eHaa8NUybpq3BgGDoTzzgtX91OmwKZN0fOqq8O7itWr4cYbw/NEGpp+zKRkvfMO/PSnYYglTrt28NnPhkKfK2bh9fr2DVf3TzwRP1Q0f364EfylL0GLFrn7+iJx1OtGStKSJfD979de5M8/H771rdwW+XSNG8Pw4fCf/wmdO8ef8/bb8L3vhUVZIg1JhV5KTlVVmNq4d2/0WLNm8Pd/H4ZNkpgB06ED3HFHGJOP8/77cM89oc2CSENRoZeSsnAh/PrX8atb27QJRfe885LNqUkTuP56uOkmaBTzG7dtW3j3sWVLsnlJ+VChl5Lx2mvwq1+F6YyZunQJRT6fC5cGDoSvfjV+uuaBYq9hHGkIKvRSEpYuhV/+Mr7I9+kDX/96WPiUb336wDe+AUcfHT324Yeh2O/YkXxeUtpU6KXorV0L999fe5H/ylfir6LzpXNn+NrX4ov91q1hOuiePcnnJaVLhV6K2pYtYQplXGE8UOQLsXXwCSfAbbdB69bRY+++C/fdF7pqiuSCCr0UrU8+gR//OH6o4+ST4ctfLswif0DHjuHKvlWr6LFly+CPf8y+ZbLIoajQS1GqqQlj8hs3Ro916hSKfNOmyed1pDp2hIkT4/8gvfQSPPts8jlJ6VGhl6L06KOwfHk03qYN/NM/wVFHJZ9TXfXoEZqdxfW+eeghWLEi+ZyktKjQS9GpqoKnnorGjzoqTF9s0yb5nOrr9NNh/PhovKYm3Gj+8MPkc5LSoUIvReW99+B3v4vGzcJVcTFv8DFoUOhbn+njj+HnPz98e2WR2qjQS9HYsyfMRolrbTB6NPTrl3xOuTZuHPTuHY2/+24YxhGpCxV6KRqTJsXffD3zzLA/aymoqIBbbokffpo7FxYtSjwlKQEq9FIUXn017BCV6cQTQw+ZUtrE4+ij4dZb43vV//73apMgR06FXgrepk1hTnmmpk3D1W8p7sPavXvYDCXTrl2haVvcKmCR2mRV6M1smJmtMLNVZnZHzPFbzewNM1tsZi+aWd9UvLuZfZqKLzazX+T6PyClraYmFLa4la/jx4d56KVq0KAwGyfT6tUwY0bi6UgRO2yhN7MK4F5gONAXmHCgkKd50N1Pc/czgO8BP0g7ttrdz0h93JqrxKU8PPlk2Ckq09lnh81DSplZ6JsfN17/xBOhx49INrK5oj8HWOXua9x9LzAJGJV+grunL0JvCWjhttTbunUwfXo03q4d3HBDaY3L16ZlS/jCF6L/15oa+O1vNeVSspNNoe8ErEt7vD4VO4iZfcXMVhOu6L+adqiHmS0ys+fM7KK4L2Bmt5hZlZlVbdZWO0LYRPs3v4k29jKDm28urG6UDe2kk8K2hJk2bIBp05LPR4pPNoU+7ropcsXu7ve6ey/gduCbqfD7QFd3HwDcBjxoZpHmrO5+v7tXuntl+/bts89eStb06WFxVKahQ6Fnz+TzybcRI+L3np09O4zZixxKNoV+PdAl7XFnYMMhzp8EjAZw9z3uvjX1+UJgNXBy3VKVcrFuHcyaFY136gRXX518PoWgceOw121FxcFx9zDlUkM4cijZFPoFQG8z62FmTYHxwEFvGM0sfS3fCGBlKt4+dTMXM+sJ9AbW5CJxKU01NaFwZU4frKgIhS5ubnm56NwZRo6Mxj/4AGbOTD4fKR6HLfTuXg1MBGYBy4HJ7r7MzO4yswM/dhPNbJmZLSYM0dyYil8MLDGz14GHgFvdXe2ZpFZPPx2W+2e68sqw72u5u+KK0O0y08yZ8UNdIgDmBbazQWVlpVdVVeU7DcmDTZvgrruiwxAnngh33lneV/PpNmyAu++Ovuvp3h1uvx0aaRlkWTKzhe5eGXdMPxJSENzD6tfMIm8Gn/+8iny6E0+Mn4XzzjswZ07i6UgRUKGXglBVBW+9FY1fdln8UEW5u/LKsO9spmnT4KOPks9HCpsKveTdp5/C5MnReNu2MGpUNC7hHc7nPx+N796tdsYSpUIveff44/EbfE+YUJoNy3KlVy+4+OJofMGC+HdHUr5U6CWv1q2L3wB7wAA47bTk8yk2o0dDq1bR+IMPhtXFIqBCL3nkHgpS5sSvpk3huuvyk1OxadkSxo6NxjduDKtmRUCFXvLo1VdhTczyuauuguOOSz6fYjVwYOiHk2nGDNi2Lfl8pPCo0Ete7N4NDz8cjXfsCIMHJ59PMTOD66+Pzp/fuxceeSQ/OUlhUaGXvJgxI/4G7PjxmjNfF506waWXRuPz56vpmajQSx5s2hRaHWQ680w45ZTk8ykVV10FrVtH43/6k7YeLHcq9JK4KVOifeabNInfI1Wy16JF/LqDdevgpZeSz0cKhwq9JGr5cliyJBq/4oqwQErq54IL4pu/TZ0a7otIeVKhl8TU1ISr+Uxt2oQNRaT+GjUK9zky7dwZ9t+V8qRCL4mZNy++le7YsVoBm0snnRQ2T880ezZs3Zp8PpJ/KvSSiN27w/BBpp49oTK2sarUxzXXRGcvVVfDo4/mJx/JLxV6ScSsWfHTKa+9NswDl9xq2xaGDInGFyyIX6QmpU2FXhrctm3xy/ErK0NjLmkYw4fHT7d86KFo2wkpbSr00uCmTYtuKNK4cRhekIbTvHn8HrOrV8PixcnnI/mTVaE3s2FmtsLMVpnZHTHHbzWzN8xssZm9aGZ90479e+p5K8xMcyvKzHvvwcsvR+OXXQbt2iWfT7m58MKwI1WmRx+NrmWQ0nXYQm9mFcC9wHCgLzAhvZCnPOjup7n7GcD3gB+kntsXGA/0A4YBP0u9npSJRx6JDhO0bBm/FZ7kXqNGMGZMNL5xoxZRlZNsrujPAVa5+xp33wtMAg5af+fu6bfZWgIHfrVHAZPcfY+7/wVYlXo9KQMrVsDSpdH4lVeGVZySjFNPhZNPjsYffxz27Ek+H0leNoW+E7Au7fH6VOwgZvYVM1tNuKL/6hE+9xYzqzKzqs2bN2ebuxQw9/julG3bwqBBiadT1szie9bv2KGe9eUim0IfN/ktcs/e3e91917A7cA3j/C597t7pbtXtm/fPouUpNAtXAhr10bjo0erO2U+dO8ev17hqafCqlkpbdkU+vVAeveMzsCGQ5w/CRhdx+dKCdi/P35xVNeu8Ss2JRmjR0NFxh2yPXtg5sz85CPJyabQLwB6m1kPM2tKuLk6Lf0EM+ud9nAEsDL1+TRgvJk1M7MeQG9gfv3TlkI2b15oRZxpzBgtjsqn9u3jNxN/7jm1Rih1hy307l4NTARmAcuBye6+zMzuMrMDs3QnmtkyM1sM3AbcmHruMmAy8CbwJPAVd9ekrhK2d2+4yZfplFOgT5/k85GDjRgR7StUXR3WOkjpymq01N1nADMyYt9K+/yfD/Hc7wDfqWuCUlzmzIHt26NxLY4qDK1bw+WXw/TpB8dffTV0EI2bcy/FTytjJWd27YpvhXvmmeFmoBSGyy8PaxnSuavhWSlToZeceeqpUOzTNWoUbgJK4WjePKxlyLRkiRqelSoVesmJHTvgmWei8fPPhw4dks9HDu2SS+C446LxuNlSUvxU6CUnZs4MN2LTNW4cNqyWwtOkCVx9dTT+1lvhQ0qLCr3U24cfwvPPR+ODBoVtAqUwnXde/Lutxx5TG+NSo0Iv9TZ9epiil65ZMxg2LD/5SHYaNYpvY/yXv8Rv4C7FS4Ve6mXjxvg2xIMHx296IYXlrLOgS5dofOpUXdWXEhV6qZfHH4eamoNjLVqEKXxS+Mxg1Kho/L33oKoq+XykYajQS51t2BBfDIYNUxviYnLqqfFbOsb9EZfipEIvdTZtWvTt/dFHqw1xsTGLX+uwcWNYMSvFT4Ve6mTtWli0KBofNizaS0UK38knh35EmeJutEvxUaGXOolrgnXssfHdEaU4xM3A2bIldCOV4qZCL0dszZr4LQJHjAgLcaQ49eoVxuszPfEE7NuXfD6SOyr0csTilsm3axfaHUhxi7uq/+gjeOGF5HOR3FGhlyPy9tvxS+RHjNAWgaWgWzcYMCAaj2txIcVDhV6y5h4/Nt+hQ1hOL6Vh5MjoTmA7dsDcuXlJR3JAhV6y9tZbsHJlNH7VVWE5vZSGE08MK2YzzZoFu3cnn4/UX1a/nmY2zMxWmNkqM7sj5vhtZvammS0xs2fMrFvasf1mtjj1oQ3LipR7/Nh8x45QWZl8PtKwrr46elX/8cdhBzEpPoct9GZWAdwLDAf6AhPMrG/GaYuASnfvDzwEfC/t2KfufkbqI+ZWjxSDpUtDs6tMI0fqar4UnXACnHtuNP7UU/Dpp8nnI/WTza/oOcAqd1/j7nuBScBB3THcfY67H9hb6BWgc27TlHyqbWy+c+f4G3dSGkaMiP4R37UrfoMZKWzZFPpOwLq0x+tTsdrcDMxMe9zczKrM7BUzi91UzsxuSZ1TtXnz5ixSkiQtWQLvvhuNx920k9Jx/PEwcGA0Pnt2dMtIKWzZFPq4X+XYBqZmdgNQCdyTFu7q7pXA9cD/mlmkfZK73+/ule5e2b59+yxSkqTUdjXfrRv07598PpKsESOgouLg2O7dodhL8cim0K8H0jtWdwY2ZJ5kZkOAO4GR7r7nQNzdN6T+XQPMBfRmv4gsWgTr10fjupovD23bwgUXROPPPBNuzkpxyKbQLwB6m1kPM2sKjAcOusYzswHAfYQivykt3sbMmqU+bwdcALyZq+SlYdXUxF/N9+wJ/foln4/kx5VXRhfD7dkTbsxKcThsoXf3amAiMAtYDkx292VmdpeZHZhFcw/QCpiSMY2yD1BlZq8Dc4DvursKfZFYuBDefz8a19V8eWnTBi66KBqfMycspJLCl9WidXefAczIiH0r7fMhtTxvHnBafRKU/KipCRtPZOrdO76drZS24cPhxRcPbm62d29YRDVuXP7ykuxoBrTEmj8/bDyRSVfz5emYY+CSS6Lx554LTc+ksKnQS8T+/WHDiUynnBI2qJDyNGwYNG16cGzfPnjyyfzkI9lToZeIV16BuOUMcS1spXy0bg2XXhqNv/ACbNuWfD6SPRV6OUh1ddhoIlO/fvEbSEt5ueKK6FaR1dUwY0b8+VIYVOjlIPPmwdat0biu5gWgVSsYPDgaf+ml+J8bKQwq9PJX+/bFX8337w/duyeejhSoyy+H5s0Pju3fH/+zI4VBhV7+6oUX4mdQXH118rlI4WrRIhT7TC+/DJs2ReOSfyr0AoQ50TNnRuMDBkDXrsnnI4Vt8OBQ8NPV1MTP1pL8U6EXIGwTl7nK0UxX8xLvqKPCjdlM8+fHr6aW/FKhF3bvDiscM511FnQ6VENqKWuXXRamXKZz11V9IVKhF559NtqJUFfzcjjNmsHQodF4VVV8x1PJHxX6MrdrV3xv8XPPDdvJiRzKJZeE9giZ4rqeSv6o0Je5uN2CGjWCq67KTz5SXJo2DQ3PMr3+OrzzTuLpSC1U6MvYzp3x+3+efz5ooy/J1kUXhVbGmaZOTT4XiadCX8ZmzQobSKRr3DhsHyeSrcaN498BvvkmrFyZfD4SpUJfpj76KEypzHTRRXDccYmnI0Vu4MD4d4FTp4aZOJJfKvRlaubMgzeRAGjSJGwbJ3KkKiriZ2mtXBmu7CW/VOjL0JYtod1Bpssug6OPTj4fKQ1nnw0dO0bjuqrPv6wKvZkNM7MVZrbKzO6IOX6bmb1pZkvM7Bkz65Z27EYzW5n6uDGXyUvdPP54aEKVrnnz+DnRItlq1AhGjYrG166FRYuSz0f+5rCF3swqgHuB4UBfYIKZ9c04bRFQ6e79gYeA76WeexzwbeBc4Bzg22YWc39ekrJhA7z6ajQ+ZAi0bJl8PlJazjgDunWLxqdODb1wJD+yuaI/B1jl7mvcfS8wCTjo77a7z3H3A7OxXwE6pz4fCsx29w/dfRswGxiWm9SlLqZNi76NbtkyvhuhyJEyi7+q/+CD+AsMSUY2hb4TsC7t8fpUrDY3Awf6IGb1XDO7xcyqzKxqc9wedpIT77wT/xZ6+PBof3GRuurbF3r3jsYffzzsRiXJy6bQW0ws9taKmd0AVAL3HMlz3f1+d69098r2WqnTYB57LBo79lgYNCjxVKSEmcHo0dH41q3xkwCk4WVT6NcDXdIedwY2ZJ5kZkOAO4GR7r7nSJ4rDe+tt2D58mh8xIgwrVIkl046CU49NRp/4onoIj1peNkU+gVAbzPrYWZNgfHAQS2LzGwAcB+hyKfvMTMLuMLM2qRuwl6RikmC3OHRR6Px9u3hgguSz0fKQ9xVfW1tN6RhHbbQu3s1MJFQoJcDk919mZndZWYHtoy+B2gFTDGzxWY2LfXcD4G7CX8sFgB3pWKSoMWL4xtMjRoVFrqINIQuXcLc+kyzZkXbYkvDapzNSe4+A5iREftW2udDDvHcB4AH6pqg1E9NTfzYfJcuUFmZfD5SXkaOhIULD55auXt3WJk9blz+8io3Whlb4l5+OUxty3TNNeGmmUhDOv740D8p09y58KHe2ydGhb6E7dsXprRlOvnkMAVOJAkjRoS+9emqq+N/NqVhqNCXsGefhW3bonFdzUuSjjkGBg+Oxl9+Gd57L/l8ypEKfYn65JMwDpppwADo2TP5fKS8DR0abbHhDo88kp98yo0KfYmaMQM+/fTgWKNG8VPeRBraUUfFt8BeuhRWrEg+n3KjQl+CtmyBOXOi8Qsv1Ibfkj+DBkHbttH4ww+rjXFDU6EvQVOnRtsQN22qDb8lvxo3rr2NcVVV8vmUExX6ErN2LcyfH41ffnm4KSaST+ecE9ZwZHrsMTU8a0gq9CXEHaZMicZbt4Yrrkg+H5FMZjB2bDRe23Cj5IYKfQl5/fWwR2emESPUhlgKR58+0K9fND5jRpgtJrmnQl8i9u8PN7UydegAF1+cfD4ihzJ2bHQtx65dobul5J4KfYl4/nnYtCkaHztWjcuk8HTqFN85dc6c+J9jqR8V+hKwa1ftrQ76908+H5FsjBoFzZodHKupiX9nKvWjQl8Cpk+Pjm2ahe6AanUgheroo8OK2UyLF2sRVa6p0Be5Dz6In61w7rnQtWvy+YgciSFDwnaWmSZPPri1sdSPCn2RmzIl+gvRtGloXCZS6Jo1i/9ZXb8eXnwx+XxKlQp9EVu2LPQKyTRsWPxVkkghOvdc6N49Gp86NdqvSeomq0JvZsPMbIWZrTKzO2KOX2xmr5lZtZldm3Fsf2p7wb9uMSj1t39//OKo447T4igpLmbw2c9G4x9/HO4/Sf0dttCbWQVwLzAc6AtMMLPMbSveBW4CHox5iU/d/YzUx8iY41IHc+fC++9H42PHQpMmiacjUi89e4b2CJmefTZ+hzQ5Mtlc0Z8DrHL3Ne6+F5gEHNSayN3fcfclgG6fJGDHDpgW897opJPgrLOSz0ckF8aMiV6k1NTApEnqbllf2RT6TsC6tMfrU7FsNTezKjN7xcxiu6Gb2S2pc6o2b958BC9dnh59NGywnM4MrrtO0ymleLVpE+4vZVq+PEy5lLrLptDHlY4j+fva1d0rgeuB/zWzXpEXc7/f3SvdvbJ9+/ZH8NLlZ80amDcvGr/wQujWLfl8RHJp6ND4nvVTpsDevcnnUyqyKfTrgfTGop2BDdl+AXffkPp3DTAXGHAE+UmaA29jM7VooZ2jpDQ0aRJ/Y3brVnjyyeTzKRXZFPoFQG8z62FmTYHxQFazZ8ysjZk1S33eDrgAeLOuyZa7F14I/eYzjRoFrVoln49IQ+jfP7675axZ6oNTV4ct9O5eDUwEZgHLgcnuvszM7jKzkQBmdraZrQfGAfeZ2bLU0/sAVWb2OjAH+K67q9DXwY4dYWw+U+fO6k4ppeXAdMvMZnzV1boxW1eNsznJ3WcAMzJi30r7fAFhSCfzefOA0+qZoxAaPcUtHpkwIWz6LVJKOnQI7RFmzTo4vmwZvPaaZpcdKZWIIvD22/DKK9H4+eeHKZUipWjEiDATJ9PkydFZZ3JoKvQFrroaHoxZhtaiRZh3LFKqmjWD8eOj8Y8+im/LLbVToS9wTz0VvwJ2zJiwF6xIKTv99Pg9FZ55Bt59N/l8ipUKfQHbuDF+a7WePcO8eZFSZxau6jNXzLrD73+vVsbZUqEvUO7whz+EoZt0jRrB3/2dVsBK+WjbNozXZ1q3Dp5+Ovl8ipEKfYF68UVYuTIav/zyMKVSpJxccUXYZzbTtGmwZUvy+RQbFfoCtH17/L6Z7drBVVcln49IvlVUwOc+F30nu28f/N//aW794ajQFxh3+OMf4+fM33BD2D1KpBz16AGXXhqNL18e3/9J/kaFvsAsWACvvx6Nn3ce9OmTfD4ihWTUqNrn1m/blnw+xUKFvoDs2BHftKx1axg3Lvl8RApN8+bhnW2m3bs1hHMoKvQFwj0sjPrkk+ix669X0zKRA049FQYOjMaXLo1fQS4q9AWjqgoWLYrGzzoLzjwz+XxECtl118Exx0TjGsKJp0JfAD76KL7NQatWoWmZiBysRYv4IZxdu+B3v9MQTiYV+jxzh9/+NvyAZpowQW0ORGrTvz+ce240vnw5zJ2beDoFTYU+z+bODT+Ymc46S61YRQ7ns5+FY4+Nxh9+OL5HVLlSoc+jDz6IXxh1zDFqcyCSjZYt4fOfj8b37YPf/CbaQqRcqdDnSXU1/OpX4Qcy0403hh9gETm8fv1g0KBofO3a0CJBVOjz5pFHQlOmTJdcEr9fpojUbswYOP74aHzWrPih0XKTVaE3s2FmtsLMVpnZHTHHLzaz18ys2syuzTh2o5mtTH3cmKvEi9kbb4R+2pmOPx7Gjk0+H5Fi16wZ/MM/xG+r+cADsHNn8jkVksMWejOrAO4FhgN9gQlm1jfjtHeBm4AHM557HPBt4FzgHODbZhazgLl8bN8eZtlkqqiAL34x/MCKyJHr0QNGjozGd+wIv3PlPOUymyv6c4BV7r7G3fcCk4BR6Se4+zvuvgTI3AZgKDDb3T90923AbGBYDvIuSjU1YVz+44+jx8aMga5dk89JpJQMHQqf+Uw0vnQpzJ6dfD6FIptC3wlIH01en4plI6vnmtktZlZlZlWbN2/O8qWLz7RpYaPvTP36weDByecjUmoaNQpDOHGTGR59NH6Ph3KQTaGPm+SX7ZugrJ7r7ve7e6W7V7Zv3z7Lly4ub7wBM2dG40cfDTfdpKmUIrly7LHhdypTTQ388pdhKKfcZFPo1wNd0h53BjZk+fr1eW7J2Lo13BDKZAZf+EIo9iKSO/37h93YMm3fDr/+dfntNZtNoV8A9DazHmbWFBgPZDs7dRZwhZm1Sd2EvSIVKxt798LPfx7f4mDUqPjxRBGpv2uugV69ovG33oLHHks+n3w6bKF392pgIqFALwcmu/syM7vLzEYCmNnZZrYeGAfcZ2bLUs/9ELib8MdiAXBXKlYW3EOP7Lj58qeeCsPK9ra0SMOrqIBbbonvFzVrVtjkp1yYF9ico8rKSq+qqsp3Gjnx9NMwZUo0ftxx8M1vavWrSBKWL4cf/Sg6vbJJE7j9dujSJf55xcbMFrp7ZdwxrYxtIMuXw0MPReNNmsCtt6rIiySlTx8YPToa37cvDKuWw2IqFfoG8P77cN998Qs0brgBunVLPieRcjZ0KFTGXOtu3RqKfVzPqVKiQp9jO3fCT34Cn34aPTZkSNjkW0SSZRa6XHbuHD22enXpb1aiQp9D+/bBz34WrhIynXKK+tiI5FOzZvClL8UPmy5YANOnJ59TUlToc6SmJsyVX7Mmet9Rv68AAAmUSURBVKxDh3D3P67hkogkp127UOwrKqLHpk+HefOSzykJKj054A5//jO89lr0WMuWMHGibr6KFIreveM3KwH4wx/CKvZSo0KfAzNnxu9RWVERrh7i+mSLSP6cdx5ceWU0XlMTJlLEvTMvZir09fTiizB1avyxG28MVw8iUnhGjoSzz47G9+2Dn/60tPacVaGvh/nzw8rXOGPHxu9QLyKFwSw0PzvllOixTz6BH/4QNm1KPK0GoUJfR4sWhc2H46ZkDRkS31BJRApL48ZheDVudez27aHYx82iKzYq9HXwxhuh3WlcB7xzz4Vrr1XbYZFi0bw5fPWrYUZOpg8/DMX+o4+SzyuXVOiP0Ouvh5V0+/dHj51+ehiXV5EXKS5HHw3/+q+hl32mzZvhf/4nFP1ipUJ/BF57DX7xi/gi37dv2PM1bn6uiBS+du3gttvi94c4UOy3bEk+r1xQoc/S/Pm1D9f07h3G+Zo0ST4vEcmdDh3ClX3cupetW0Ox37gx+bzqS4U+C88+W/uuNCedFBZENW2afF4iknsnnlh7sd+2De65B9auTT6v+lChPwT3MEf+z3+OP/6Zz4SbOM2bJ5uXiDSsLl3ga1+L37Rk5074/vdDK/JioUJfi/37w3LoGTPij/ftG67kmzVLNi8RSUanTqHYx43Z79kTutTOn598XnWRVaE3s2FmtsLMVpnZHTHHm5nZn1PHXzWz7ql4dzP71MwWpz5+kdv0G8auXfDjH8NLL8UfP+MM+PKXNVwjUuo6doRvfAPato0e278/DOk+/njhtzg+bKE3swrgXmA40BeYYGZ9M067Gdjm7icBPwT+X9qx1e5+Rurj1hzl3WA2b4bvfjdsIBznwgvhH/9RN15FysXxx8O//Vu4wo8zfXoo+IW8eUk2V/TnAKvcfY277wUmAaMyzhkF/C71+UPAYLPim02+bBn893/Xflf9yivDDlFqNyxSXo49Fr7+9TD5Is6CBeEmbaHOtc+mZHUC1qU9Xp+KxZ7j7tXAduDAm50eZrbIzJ4zs4vqmW+DcA8dKH/ykzBsk6lRI7j+ehg1SouhRMpVixbwL/8CZ50Vf3ztWvjOd2DFimTzykY2hT6utGWOSNV2zvtAV3cfANwGPGhmkVsbZnaLmVWZWdXmzZuzSCl3PvkkrHR97LH4cbbmzcNN10suSTQtESlATZqEhZFxLY4BPv44tEx48snCGrfPptCvB9Jb/nQGNtR2jpk1Bo4BPnT3Pe6+FcDdFwKrgZMzv4C73+/ule5e2b59+yP/X9TRypVw992hrUGctm3h9tuhX7/EUhKRAmcW3t3fdFP8Snh3ePRR+NGPYMeOxNOLlU2hXwD0NrMeZtYUGA9MyzhnGnBj6vNrgWfd3c2sfepmLmbWE+gN5L2l//794U75978fFkDE6dMH7rwzLJ4QEck0cGAYt4/rjwNhnv1dd8HSpcnmFeewhT415j4RmAUsBya7+zIzu8vMRqZO+zXQ1sxWEYZoDkzBvBhYYmavE27S3urueb1dsWFDmFUzfXrtb62GDw8LobT9n4gcSs+e4YLw5Mg4RbBzZ7j39/vfw6efJptbOvNCGkgCKisrvaqqKuevu38/PPVUKPDV1fHntGgRuk+ecUbOv7yIlLD9+8Mq+lmzaj+nTRv43OcabijYzBa6e2XcscYN8yULy5o1YZXrhsw7C2lOOgluvhmOOy65vESkNFRUwJgxYbeqBx4IV/KZtm0LCzHPPhuuuy5+xW1DKekr+p07w1/ZF16o/ZxGjcId9BEjND9eROpvxw747W/DupzaHHVUuKF7ySW5qzuHuqIvyUJfXQ1z5oRhmt27az+vY8dw57x793p9ORGRg7jDiy/ClCmhL05tOnaEceNyM5xTNoXePaxQmzYttDKojVnY03XkSLUyEJGGs3VruBFbW0uVA/r2hdGjoVu3un+tki/07mEf18ceg/feO/S53bqFNgZdu9YjSRGRLLnDq6/C5MlhgeahnHlmGNI54YQj/zolfTN2zRp46CFYvfrQ5zVvHq7gL71UY/EikhwzOO88OPVUePhhmDev9nNfew0WLYLzzw8F/5hjcpND0Ze8Dz44dJE3g4sugv/6Lxg8WEVeRPKjVaswffs//gN69ar9vAPvAOJ2tKuror+iP++8MHf1gw+ix/r2hbFjoXPn5PMSEYnTrVvocV9VFYab4zYcHzQozLvPlaIv9I0ahbc49933t1jPnnDNNbWvVhMRySezMJ9+wICwwdETT8D27eFY8+ZhdX4uFX2hh/DN6tYtTKscPRpOO03thEWk8DVuHObSDxwYpoQ/+WQYYm7VKrdfpyRm3UD4a9i6tcbgRaR47doVVtnWZS/qkp51c0Cu7k6LiORLixYN87q6/hURKXEq9CIiJU6FXkSkxKnQi4iUOBV6EZESV3DTK81sM7A233lkaAfErF8rK/oeBPo+6HtwQKF9H7q5e/u4AwVX6AuRmVXVNj+1XOh7EOj7oO/BAcX0fdDQjYhIiVOhFxEpcSr02bk/3wkUAH0PAn0f9D04oGi+DxqjFxEpcbqiFxEpcSr0IiIlToU+S2Z2j5m9ZWZLzOxRMzs23zklzczGmdkyM6sxs6KYVpYrZjbMzFaY2SozuyPf+eSDmT1gZpvMbGm+c8kXM+tiZnPMbHnqd+Gf851TNlToszcbONXd+wNvA/+e53zyYSkwBng+34kkycwqgHuB4UBfYIKZ9c1vVnnxW2BYvpPIs2rga+7eBzgP+Eox/Cyo0GfJ3Z9y9+rUw1eAstuJ1t2Xu/uKfOeRB+cAq9x9jbvvBSYBo/KcU+Lc/Xngw3znkU/u/r67v5b6fCewHOiU36wOT4W+bv4BmJnvJCQxnYB1aY/XUwS/3NKwzKw7MAB4Nb+ZHF7J7DCVC2b2NHBCzKE73X1q6pw7CW/f/phkbknJ5ntQhuJ2INa85DJmZq2Ah4F/cfcd+c7ncFTo07j7kEMdN7MbgauAwV6iCxAO9z0oU+uBLmmPOwMb8pSL5JmZNSEU+T+6+yP5zicbGrrJkpkNA24HRrr7rnznI4laAPQ2sx5m1hQYD0zLc06SB2ZmwK+B5e7+g3znky0V+uz9FGgNzDazxWb2i3wnlDQzu8bM1gMDgSfMbFa+c0pC6ib8RGAW4ebbZHdflt+skmdmfwJeBj5jZuvN7OZ855QHFwCfAy5L1YHFZnZlvpM6HLVAEBEpcbqiFxEpcSr0IiIlToVeRKTEqdCLiJQ4FXoRkRKnQi8iUuJU6EVEStz/B7d/fADztCq7AAAAAElFTkSuQmCC\n",
      "text/plain": [
       "<Figure size 432x288 with 1 Axes>"
      ]
     },
     "metadata": {
      "needs_background": "light"
     },
     "output_type": "display_data"
    }
   ],
   "source": [
    "import numpy as np\n",
    "from scipy.stats import norm\n",
    "import matplotlib.pyplot as plt\n",
    "r=norm.rvs(loc=0,scale=1,size=1000)#正态分布\n",
    "x=np.linspace(norm.ppf(0.01),\n",
    "             norm.ppf(0.99),100)\n",
    "fig,ax=plt.subplots(1,1)\n",
    "ax.plot(x,norm.pdf(x),\n",
    "       'blue',lw=5,alpha=0.6,label='norm pdf')\n",
    "plt.show()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 1,
   "metadata": {
    "scrolled": true
   },
   "outputs": [
    {
     "ename": "FileNotFoundError",
     "evalue": "[Errno 2] No such file or directory: 'data/BABAnews.txt'",
     "output_type": "error",
     "traceback": [
      "\u001b[1;31m---------------------------------------------------------------------------\u001b[0m",
      "\u001b[1;31mFileNotFoundError\u001b[0m                         Traceback (most recent call last)",
      "\u001b[1;32m<ipython-input-1-7764569470ca>\u001b[0m in \u001b[0;36m<module>\u001b[1;34m\u001b[0m\n\u001b[1;32m----> 1\u001b[1;33m \u001b[1;32mwith\u001b[0m \u001b[0mopen\u001b[0m\u001b[1;33m(\u001b[0m\u001b[1;34m\"data/BABAnews.txt\"\u001b[0m\u001b[1;33m,\u001b[0m \u001b[0mmode\u001b[0m \u001b[1;33m=\u001b[0m \u001b[1;34m'r'\u001b[0m\u001b[1;33m)\u001b[0m \u001b[1;32mas\u001b[0m \u001b[0mmyfile\u001b[0m\u001b[1;33m:\u001b[0m\u001b[1;33m\u001b[0m\u001b[1;33m\u001b[0m\u001b[0m\n\u001b[0m\u001b[0;32m      2\u001b[0m         \u001b[0mprint\u001b[0m\u001b[1;33m(\u001b[0m\u001b[0mmyfile\u001b[0m\u001b[1;33m.\u001b[0m\u001b[0mread\u001b[0m\u001b[1;33m(\u001b[0m\u001b[1;33m)\u001b[0m\u001b[1;33m)\u001b[0m\u001b[1;33m\u001b[0m\u001b[1;33m\u001b[0m\u001b[0m\n",
      "\u001b[1;31mFileNotFoundError\u001b[0m: [Errno 2] No such file or directory: 'data/BABAnews.txt'"
     ]
    }
   ],
   "source": [
    "with open(\"data/BABAnews.txt\", mode = 'r') as myfile: \n",
    "        print(myfile.read())"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "with open(\"新建.txt\",mode='r') as asd:\n",
    "    print(asd.read())"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "myfile.closed"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "If you’re not using the with keyword, then you should call f.close() to close the file and immediately free up any system resources used by it. For example:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {
    "scrolled": true
   },
   "outputs": [],
   "source": [
    "myfile = open(\"data/BABAnews.txt\", mode = 'r')\n",
    "text = myfile.read()\n",
    "myfile.close()\n",
    "print(text)"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {
    "scrolled": true
   },
   "outputs": [],
   "source": [
    "myfile.read() # can be read after the file being closed"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "To read the contents of a file, we use `f.read(size)`. "
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "myfile = open(\"data/BABAnews.txt\", mode = 'r')\n",
    "myfile.read()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "myfile.read()\n",
    "myfile.close()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "`f.readline()` reads a single line from the file; a newline character (\\n) is left at the end of the string, and is only omitted on the last line of the file if the file doesn’t end in a newline. This makes the return value unambiguous; if `f.readline()` returns an empty string, the end of the file has been reached, while a blank line is represented by '\\n', a string containing only a single newline. \n",
    "\n",
    "If you want to read all the lines of a file in a list you can also use `list(f)` or `f.readlines()`."
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "myfile = open(\"data/BABAnews.txt\", mode = 'r')\n",
    "myfile.readline()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "myfile.readlines()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "**Note:** Both `read()` and `readlines()` come with the concept of a cursor. After either command is executed, the cursor moves to the end of the file, leaving nothing more to read in. Therefore, once a file content has been read in, another attempt to read from the file object will produce an empty data object. If for some reason you must read the file content again, you must close and re-open the file. \n",
    "\n",
    "Lastly, rather than loading the entire file content into memory, you can iterate through the file object line by line using the `for` loop. This method is more memory-efficient and therefore recommended when dealing with a very large file. "
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "with open('新建.txt', 'r') as myfile:\n",
    "    for line in myfile:\n",
    "        print(line,end='')"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "Writing methods also come in a pair: `write()` and `writelines()`. Like the corresponding reading methods, `write()` handles a single string, while `writelines()` handles a list of strings. \n",
    "\n",
    "Below, `write()` writes a single string each time to the designated output file:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "file = open('data/testfile.txt', 'a') \n",
    "file.write('Hello World \\n') \n",
    "file.write('and this is another line.') \n",
    "file.close() "
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "This time, we have `tobuy`, a list of strings, which `writelines()` writes out at once:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "tobuy = ['milk\\n', 'butter\\n', 'coffee beans\\n', 'arugula\\n']\n",
    "file = open('data/grocerylist.txt', 'w')\n",
    "file.writelines(tobuy) # writelines(list)\n",
    "file.close()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "**Only the string type can be written.** Writing methods only works with strings: `write()` takes a single string, and `writelines()` takes a list which contains strings only. Non-string type data must be first coerced into the string type by using the `str()` function. "
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "pi = 3.141592\n",
    "fout = open('data/math.txt', 'w')\n",
    "fout.write(\"Pi's value is \")\n",
    "fout.write(pi) # trying to write float, doesn't work"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "fout.write(str(pi))\n",
    "fout.close()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "### Other ways to read text files"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "We can also use other modules to read text files. For example, we can use **numpy** to read *txt* file.  **csv** and **pandas** can be used to read *csv* files."
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "import numpy as np \n",
    "data = np.loadtxt('data/BJsales.txt', skiprows = 1, delimiter='\\t')\n",
    "print(data)"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {
    "scrolled": true
   },
   "outputs": [],
   "source": [
    "import csv\n",
    "with open('data/Titanic.csv') as csvfile:\n",
    "    titanicReader = csv.reader(csvfile)\n",
    "    for row in titanicReader:\n",
    "        print('  '.join(row))"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "The `pandas` I/O API is a set of top level `reader` functions accessed like `read_csv()` that generally return a pandas object. These functions includes\n",
    "\n",
    "    read_excel\n",
    "    read_hdf\n",
    "    read_sql\n",
    "    read_json\n",
    "    read_msgpack (experimental)\n",
    "    read_html\n",
    "    read_gbq (experimental)\n",
    "    read_stata\n",
    "    read_sas\n",
    "    read_clipboard\n",
    "    read_pickle\n",
    "    \n",
    "See [pandas IO tools](http://pandas.pydata.org/pandas-docs/stable/io.html) for detailed explanation."
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {
    "scrolled": false
   },
   "outputs": [],
   "source": [
    "import pandas as pd\n",
    "titanicData = pd.read_csv('data/Titanic.csv') \n",
    "titanicData.head()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "We could also have used `read_table()` and specifying the delimiter:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "titanicData = pd.read_table('data/Titanic.csv', sep=',')\n",
    "titanicData.head()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "A file will not always have a header row. To read this in, you have a couple of options. You can allow pandas to assign default column names, or you can specify names yourself:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "sales = pd.read_csv('data/BJsales.csv', header = None)\n",
    "sales.head()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "sales = pd.read_csv('data/BJsales.csv', names=['ID', 'Time', 'Value'])\n",
    "sales.head()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "Data can also be exported to delimited format. Let’s consider one of the CSV files read above:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "sales.to_csv('data/sales.csv', index = False)"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "## Reading a Microsoft Excel file using `pandas`"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "Excel files are a huge part of any business operation and it becomes imperative that you learn exactly how to import these into python for data analysis. In order to do this we can use the code snippet shown below:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "import pandas as pd\n",
    "file = 'data/example.xlsx'\n",
    "myfile = pd.ExcelFile(file) \n",
    "print(myfile.sheet_names) #printing out all the sheet names in the excel file\n",
    "dataframe = myfile.parse('HairEyeColor') #extracting data from the first sheet as a dataframe\n",
    "dataframe.head()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "## Reading a pdf file using `PyPDF2`\n",
    "\n",
    "`PyPDF2` is a python library built as a PDF toolkit. It is capable of:\n",
    "\n",
    "- Extracting document information (title, author, …)\n",
    "- Splitting documents page by page\n",
    "- Merging documents page by page\n",
    "- Cropping pages\n",
    "- Merging multiple pages into a single page\n",
    "- Encrypting and decrypting PDF files\n",
    "- and more!\n",
    "\n",
    "Here we only demonstrate how to read a pdf file. Please find more details on https://pypi.org/project/PyPDF2/."
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "import PyPDF2\n",
    "pdfFileObj = open('data/Li2011Wiley.pdf', 'rb')\n",
    "pdfReader = PyPDF2.PdfFileReader(pdfFileObj)\n",
    "pdfReader.numPages\n",
    "pageObj = pdfReader.getPage(0)\n",
    "print(pageObj.extractText())\n",
    "pdfFileObj.close()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "## Reading a word file using `docx`"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "import docx\n",
    "doc = docx.Document('data/Li2011Wiley.docx')\n",
    "print(len(doc.paragraphs))\n",
    "print(doc.paragraphs[0].text)"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "## Interacting with HTML and Web APIs"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "Many websites have public APIs providing data feeds via JSON or some other format. There are a number of ways to access these APIs from Python; one easy-to-use method that I recommend is the requests package (http://docs.python-requests.org). "
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "import requests\n",
    "url = 'https://sp0.baidu.com/8aQDcjqpAAV3otqbppnN2DJv/api.php?resource_id=6899&query=失信执行人名单&iname=中国银行'\n",
    "resp = requests.get(url)\n",
    "import json\n",
    "data = json.loads(resp.text, encoding='gb18030')\n",
    "data.keys()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "result = data['data'][0]['result']\n",
    "for i in range(len(result)):\n",
    "        data = result[i]\n",
    "        caseCode = data['caseCode']\n",
    "        areaName = data['areaName']\n",
    "        businessEntity = data['businessEntity']\n",
    "        courtName = data['courtName']\n",
    "        duty = data['duty']\n",
    "        performance = data['performance']\n",
    "        disruptTypeName = data['disruptTypeName']\n",
    "        publishDate = data['publishDate']\n",
    "        regDate = data['regDate']\n",
    "        gistId = data['gistId']\n",
    "        gistUnit = data['gistUnit']\n",
    "        cardNum = data['cardNum']\n",
    "        print('  '.join([caseCode, areaName, businessEntity, courtName, duty, performance, disruptTypeName, publishDate, regDate, gistId, gistUnit, cardNum]))\n"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "## Interacting with Databases"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "In many applications data rarely comes from text files, that being a fairly inefficient way to store large amounts of data. SQL-based relational databases (such as SQL Server, PostgreSQL, and MySQL) are in wide use, and many alternative non-SQL (so-called NoSQL) databases have become quite popular. The choice of database is usually de- pendent on the performance, data integrity, and scalability needs of an application.\n",
    "\n",
    "Loading data from SQL into a DataFrame is fairly straightforward, and pandas has some functions to simplify the process. As an example, I’ll use an in-memory SQLite database using Python’s built-in sqlite3 driver. Here’s a short Python program that selects latitudes and longitudes from an SQLite database stored in a file called survey.db:"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": [
    "import sqlite3\n",
    "connection = sqlite3.connect(\"survey.db\")\n",
    "cursor = connection.cursor()\n",
    "cursor.execute(\"SELECT Site.lat, Site.long FROM Site;\")\n",
    "results = cursor.fetchall()\n",
    "for r in results:\n",
    "    print(r)\n",
    "cursor.close()\n",
    "connection.close()"
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "The program starts by importing the sqlite3 library. If we were connecting to MySQL, DB2, or some other database, we would import a different library, but all of them provide the same functions, so that the rest of our program does not have to change (at least, not much) if we switch from one database to another.\n",
    "\n",
    "Line 2 establishes a connection to the database. Since we’re using SQLite, all we need to specify is the name of the database file. Other systems may require us to provide a username and password as well. Line 3 then uses this connection to create a cursor. Just like the cursor in an editor, its role is to keep track of where we are in the database.\n",
    "\n",
    "On line 4, we use that cursor to ask the database to execute a query for us. The query is written in SQL, and passed to cursor.execute as a string. It’s our job to make sure that SQL is properly formatted; if it isn’t, or if something goes wrong when it is being executed, the database will report an error.\n",
    "\n",
    "The database returns the results of the query to us in response to the cursor.fetchall call on line 5. This result is a list with one entry for each record in the result set; if we loop over that list (line 6) and print those list entries (line 7), we can see that each one is a tuple with one element for each field we asked for.\n",
    "\n",
    "Finally, lines 8 and 9 close our cursor and our connection, since the database can only keep a limited number of these open at one time. Since establishing a connection takes time, though, we shouldn’t open a connection, do one operation, then close the connection, only to reopen it a few microseconds later to do another operation. Instead, it’s normal to create one connection that stays open for the lifetime of the program."
   ]
  },
  {
   "cell_type": "markdown",
   "metadata": {},
   "source": [
    "# Lab\n",
    "\n",
    "Read your five different types of your own files to Python from your hard disk."
   ]
  }
 ],
 "metadata": {
  "kernelspec": {
   "display_name": "Python 3",
   "language": "python",
   "name": "python3"
  },
  "language_info": {
   "codemirror_mode": {
    "name": "ipython",
    "version": 3
   },
   "file_extension": ".py",
   "mimetype": "text/x-python",
   "name": "python",
   "nbconvert_exporter": "python",
   "pygments_lexer": "ipython3",
   "version": "3.7.6"
  }
 },
 "nbformat": 4,
 "nbformat_minor": 2
}
