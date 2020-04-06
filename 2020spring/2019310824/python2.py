with open('Advisor.txt', mode='r') as my_file:
	print(my_file.read())
	
	my_file = open("Advisor.txt", mode = 'r')
	text = my_file.read()
	my_file.close()
	print(text)
	print('\n')
	
	my_file = open("Advisor.txt", mode = 'r')
	my_file.read()
	
	my_file = open('Advisor.txt' , mode='r')
	print(my_file.readline())
	
	with open('Titanic.csv', 'r') as my_file:
	for line in my_file:
	print(line, end ='')
	
	file = open('testfile.txt', 'w') 
	file.write('Hello World \n') 
	file.write('today is a happy day.') 
	file.close() 
	
	tobuy = ['strawberry\t', 'egg\t', 'wine\t', 'apple\t']
	file = open('grocerylist.txt', 'w')
	file.writelines(tobuy) # writelines(list)
	file.close()
	
	pi = 3.141592
	fout = open('math.txt', 'w')
	fout.write("Pi's value roughly equals ")
	fout.write(str(pi))
	fout.close()
	
	import numpy as np 
	data = np.loadtxt('BJsales.txt', skiprows = 1, delimiter='\t')
	print(data)
	
	import csv
	with open('Titanic.csv') as csvfile:
	titanicReader = csv.reader(csvfile)
	for row in titanicReader:
	print(' '.join(row))
	
	import pandas as pd
	titanicData = pd.read_csv('Titanic.csv') 
	print(titanicData.head())
	
	import pandas as pd
	sales = pd.read_csv('BJsales.csv',names=['ID','TIME','VALUE'])
	print(sales.head)
	
	import pandas as pd
	file ='example.xlsx'
	my_file = pd.ExcelFile(file)
	print(my_file.sheet_names)
	dataframe = my_file.parse('HairEyeColor')
	print(dataframe.head())
	
	import PyPDF2
	pdfFileObj = open('BWF.pdf' , 'rb')
	pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
	pdfReader.numPages
	pageObj=pdfReader.getPage(0)
	print(pageObj.extractText())
	pdfFileObj.close()
	
	import docx
	doc = docx.Document('BWF.docx')
	print(len(doc.paragraphs))
	print(doc.paragraphs[0].text)
	
	import sqlite3
	connection = sqlite3.connect("survey.db")
	cursor = connection.cursor()
	cursor.execute("SELECT Site.lat, Site.long FROM Site;")
	results = cursor.fetchall()
	for r in results:
	print(r)
	cursor.close()
	connection.close()
	
	
	import MySQLdb
	db = MySQLdb.connect("survey.db" )
	cursor = db.cursor()
	cursor.execute("SELECT VERSION()")
	data = cursor.fetchone()
	print(data)
	db.close()
	
	import numpy as np
	from scipy import linalg
	A = np.array([[1,2],[3,4]])
	print(A)
	print(linalg.inv(A))
	
	b=np.array([[5],[6]])
	print(b.T)
	print(A*b)
	print(np.linalg.solve(A,b))
	print(linalg.det(A))
	
	import numpy as np
	from scipy import linalg
	import matplotlib.pyplot as plt
	
	c1, c2 = 5.0, 2.0
	i = np.r_[1:11]
	xi = 0.1*i
	yi = c1*np.exp(-xi) + c2*xi
	zi = yi + 0.05 * np.max(yi) * np.random.randn(len(yi))
	A = np.c_[np.exp(-xi)[:, np.newaxis], xi[:, np.newaxis]]
	print(linalg.lstsq(A, zi))
	
	import numpy as np
	from scipy import linalg
	A = np.array([[1,2],[3,4]])
	la,v = linalg.eig(A)
	l1,l2 = la
	print(l1, l2) 
	print(v[:,0]) 
	print(v[:,1]) 
	print(np.sum(abs(v**2),axis=0)) 
	v1 = np.array(v[:,0]).T
	print(linalg.norm(A.dot(v1)-l1*v1))
	
	from scipy.stats import norm
	r = norm.rvs(loc=0, scale=1, size=1000)
	print(r)
	print(norm.stats(moments='mvsk'))
	
	
	from scipy import stats
	import numpy as np
	x = np.random.random(10)
	y = np.random.random(10)
	slope, intercept, r_value, p_value, std_err = stats.linregress(x,y)
	print({'slope':slope,'intercept':intercept})
	print({'p_value':p_value,'r-squared':round(r_value**2,2)})
	
	import numpy as np
	from scipy.optimize import minimize
	
	Define the function
	def rosen(x):
	return sum(100.0*(x[1:]-x[:-1]**2.0)**2.0 + (1-x[:-1])**2.0)
	x0 = np.array([1.3, 0.7, 0.8, 1.9, 1.2])
	res = minimize(rosen, x0, method='nelder-mead',
	options={'xtol': 1e-8, 'disp': True})
	print(res.x)
	
	
	import statsmodels.api as sm
	import statsmodels.formula.api as smf
	star98 = sm.datasets.star98.load_pandas().data
	formula = 'SUCCESS ~ LOWINC + PERASIAN + PERBLACK + PERHISP + PCTCHRT + \
	PCTYRRND + PERMINTE*AVYRSEXP*AVSALK + PERSPENK*PTRATIO*PCTAF'
	dta = star98[['NABOVE', 'NBELOW', 'LOWINC', 'PERASIAN', 'PERBLACK', 'PERHISP',
	'PCTCHRT', 'PCTYRRND', 'PERMINTE', 'AVYRSEXP', 'AVSALK',
	'PERSPENK', 'PTRATIO', 'PCTAF']].copy()
	endog = dta['NABOVE'] / (dta['NABOVE'] + dta.pop('NBELOW'))
	del dta['NABOVE']
	dta['SUCCESS'] = endog
	mod1 = smf.glm(formula=formula, data=dta, family=sm.families.Binomial()).fit()
	print(mod1.summary())
	
	import numpy as np
	from scipy import stats
	import pandas as pd
	import matplotlib.pyplot as plt
	import statsmodels.api as sm
	from statsmodels.graphics.api import qqplot
	print(sm.datasets.sunspots.NOTE)
	
	arma_mod20 = sm.tsa.ARMA(dta, (2,0)).fit(disp=False)
	print(arma_mod20.params)
	arma_mod30 = sm.tsa.ARMA(dta, (3,0)).fit(disp=False)
	resid = arma_mod30.resid
	stats.normaltest(resid)
	fig = plt.figure(figsize=(12,8))
	ax = fig.add_subplot(111)
	fig = qqplot(resid, line='q', ax=ax, fit=True)
	predict_sunspots = arma_mod30.predict('1990', '2012', dynamic=True)
	print(predict_sunspots)
	fig, ax = plt.subplots(figsize=(12, 8))
	ax = dta.ix['1950':].plot(ax=ax)
	fig = arma_mod30.plot_predict('1990', '2012', dynamic=True, ax=ax, plot_insample=False)
